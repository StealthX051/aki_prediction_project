import logging
import math
import os
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Set, Tuple

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt, Inches
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.colors import LinearSegmentedColormap, to_hex
from mpl_toolkits.axes_grid1.inset_locator import inset_axes
import numpy as np
from joblib import Parallel, delayed
from sklearn.calibration import calibration_curve
from sklearn.metrics import (
    average_precision_score,
    precision_recall_curve,
    roc_auc_score,
    roc_curve,
)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from reporting.display_dictionary import load_display_dictionary
from reporting.feature_set_display import FeatureSetDisplay
from artifact_paths import enforce_storage_policy, get_paper_dir, get_results_dir
from reporting.manuscript_assets import (
    export_dataframe_bundle,
    export_table_bundle,
    refresh_paper_bundle,
    save_figure_bundle,
    write_docx_sections,
    write_markdown_sections,
    write_pdf_sections,
)

# Constants
EXPERIMENTS_DIR = get_results_dir(PROJECT_ROOT)
PAPER_DIR = get_paper_dir(PROJECT_ROOT)
FIGURES_DIR = Path(os.getenv("FIGURES_DIR", PAPER_DIR / "figures"))
TABLES_DIR = Path(os.getenv("TABLES_DIR", PAPER_DIR / "tables"))
REPORTS_DIR = Path(os.getenv("REPORTS_DIR", PAPER_DIR / "reports"))
# Backwards-compatible alias used by imports in reporting.make_report
RESULTS_DIR = EXPERIMENTS_DIR

# Ensure directories exist
enforce_storage_policy(
    {
        "paper_figures_dir": FIGURES_DIR,
        "paper_tables_dir": TABLES_DIR,
        "paper_reports_dir": REPORTS_DIR,
    }
)
FIGURES_DIR.mkdir(parents=True, exist_ok=True)
TABLES_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

logger = logging.getLogger(__name__)

# --- Plot configuration helpers ------------------------------------------------


def _env_bool(name: str, default: bool) -> bool:
    """Parse a boolean-like environment variable."""
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "y", "on"}


def _env_set(name: str) -> Optional[Set[str]]:
    """Parse a comma-separated env var into a set of strings."""
    raw = os.getenv(name)
    if not raw:
        return None
    items = {item.strip() for item in raw.split(",") if item.strip()}
    return items or None


@dataclass
class PlotConfig:
    """Configurable plotting defaults with environment overrides.

    Defaults mirror the current behavior (plot all models/feature sets) but
    switch calibration bins to quantile-based for stability and prefer
    calibrated probabilities when available.
    """

    n_bins: int = 10
    bin_strategy: str = "quantile"
    prefer_calibrated: bool = True
    feature_set_allowlist: Optional[Set[str]] = None
    feature_set_blocklist: Optional[Set[str]] = None
    model_allowlist: Optional[Set[str]] = None
    model_blocklist: Optional[Set[str]] = None
    show_bin_counts: bool = True
    show_prob_hist: bool = True
    show_xlim_inset: bool = True
    max_count_to_annotate: int = 30
    n_jobs: int = -2
    show_class_balance: bool = False

    @classmethod
    def from_env(cls) -> "PlotConfig":
        """Create a PlotConfig using optional environment overrides."""
        return cls(
            n_bins=int(os.getenv("CALIBRATION_N_BINS", cls.n_bins)),
            bin_strategy=os.getenv("CALIBRATION_BIN_STRATEGY", cls.bin_strategy),
            prefer_calibrated=_env_bool("PLOT_PREFER_CALIBRATED", cls.prefer_calibrated),
            feature_set_allowlist=_env_set("PLOT_FEATURE_SET_ALLOWLIST"),
            feature_set_blocklist=_env_set("PLOT_FEATURE_SET_BLOCKLIST"),
            model_allowlist=_env_set("PLOT_MODEL_ALLOWLIST"),
            model_blocklist=_env_set("PLOT_MODEL_BLOCKLIST"),
            show_bin_counts=_env_bool("CALIBRATION_SHOW_BIN_COUNTS", cls.show_bin_counts),
            show_prob_hist=_env_bool("CALIBRATION_SHOW_PROB_HIST", cls.show_prob_hist),
            show_xlim_inset=_env_bool("CALIBRATION_SHOW_XLIM_INSET", cls.show_xlim_inset),
            max_count_to_annotate=int(os.getenv("CALIBRATION_MAX_COUNT_ANNOTATE", cls.max_count_to_annotate)),
            n_jobs=int(os.getenv("PLOT_N_JOBS", cls.n_jobs)),
            show_class_balance=_env_bool("PR_SHOW_CLASS_BALANCE", cls.show_class_balance),
        )


DEFAULT_FEATURE_SET_MAPPING = {
    'preop_only': 'Preoperative Only',
    'all_waveforms': 'All Waveforms',
    'preop_and_all_waveforms': 'Preop + All Waveforms',
    'awp_only': 'Airway Pressure Only',
    'co2_only': 'CO2 Only',
    'ecg_only': 'ECG Only',
    'pleth_only': 'Plethysmography Only',
    'monitors_only': 'Standard Monitors Only',
    'ventilator_only': 'Ventilator Only',
    'preop_and_awp': 'Preop + Airway Pressure',
    'preop_and_co2': 'Preop + CO2',
    'preop_and_ecg': 'Preop + ECG',
    'preop_and_pleth': 'Preop + Plethysmography',
    'preop_and_all_minus_awp': 'Preop + All (No AWP)',
    'preop_and_all_minus_co2': 'Preop + All (No CO2)',
    'preop_and_all_minus_ecg': 'Preop + All (No ECG)',
    'preop_and_all_minus_pleth': 'Preop + All (No Pleth)',
    # Aeon / Multirocket Mappings
    'all_fused': 'All Channels (Fused)',
    'all_waveonly': 'All Channels (Waveform Only)',
}

try:
    DISPLAY_DICTIONARY = load_display_dictionary()
except FileNotFoundError:
    logger.warning(
        "Display dictionary not found; using default feature set labels only. "
        "Create metadata/display_dictionary.json to override."
    )
    DISPLAY_DICTIONARY = None
FEATURE_SET_DISPLAY = FeatureSetDisplay.from_defaults(
    display_dictionary=DISPLAY_DICTIONARY,
    default_label_map=DEFAULT_FEATURE_SET_MAPPING,
)
FEATURE_SET_MAPPING = FEATURE_SET_DISPLAY.label_map

NEURIPS_CSS = """
<style>
    body {
        font-family: 'Times New Roman', Times, serif;
        background-color: #ffffff;
        color: #000000;
        margin: 40px;
    }

    h1 {
        font-weight: bold;
        font-size: 24pt;
        margin-bottom: 10px;
        text-align: center;
    }

    .subtitle {
        font-size: 14pt;
        text-align: center;
        margin-bottom: 30px;
        font-style: italic;
    }

    table {
        border-collapse: collapse;
        width: 100%;
        margin-bottom: 40px;
        font-size: 11pt;
        border-top: 2px solid #000000;
        border-bottom: 2px solid #000000;
    }

    th {
        background-color: #ffffff;
        font-weight: bold;
        text-align: left;
        padding: 8px;
        border-bottom: 1px solid #000000;
    }

    td {
        padding: 8px;
        border-bottom: 1px solid #e0e0e0;
    }

    tr:last-child td {
        border-bottom: none;
    }

    /* Checkbox grid for feature inputs */
    .inputs-grid {
        border-collapse: collapse;
        margin: 0 auto;
    }

    .inputs-grid th,
    .inputs-grid td {
        border: 1px solid #000000;
        padding: 4px 6px;
        text-align: center;
        font-size: 10pt;
    }

    .inputs-grid th {
        background-color: #f5f5f5;
        font-weight: bold;
    }

    .inputs-grid td.checked {
        background-color: #f8fdf8;
    }

    .inputs-grid-unknown {
        font-style: italic;
        text-align: center;
        display: block;
        padding: 4px 0;
    }

    /* Tasteful highlighting */
    .highlight-high {
        background-color: #f0f7f0; /* Very light green */
        font-weight: bold;
    }

    .footer {
        font-size: 10pt;
        color: #666666;
        margin-top: 20px;
        text-align: center;
        border-top: 1px solid #000000;
        padding-top: 10px;
    }
</style>
"""

METRIC_DISPLAY_MAP = {
    'auroc': 'AUROC',
    'auprc': 'AUPRC',
    'brier': 'Brier Score',
    'sensitivity': 'Sensitivity',
    'specificity': 'Specificity',
    'f1': 'F1 Score',
}


def prepare_metrics_for_display(
    summary_df: pd.DataFrame, renderer: FeatureSetDisplay = FEATURE_SET_DISPLAY
) -> pd.DataFrame:
    """Prepare metrics_summary output for styling and reporting."""

    metrics_df = summary_df.copy()
    metrics_df["Feature Set Key"] = metrics_df["feature_set"]

    rename_map = {
        'outcome': 'Outcome',
        'branch': 'Branch',
        'feature_set': 'Feature Set',
        'model_name': 'Model',
    }
    metrics_df.rename(columns=rename_map, inplace=True)

    # Populate display-friendly feature set variants
    metrics_df["Feature Set Label"] = metrics_df["Feature Set Key"].apply(renderer.as_label)
    metrics_df["Feature Set"] = metrics_df["Feature Set Label"]
    # Expand checkbox columns if enabled
    if renderer.show_checkbox:
        for comp in renderer.config.components:
            col_name = comp.label
            def _mark(fs: str, comp_key: str = comp.key) -> str:
                flags, known = renderer.component_flags(fs)
                if not known:
                    return "?"
                return "✓" if flags.get(comp_key, False) else ""
            metrics_df[col_name] = metrics_df["Feature Set Key"].apply(_mark)
    # Text fallback for DOCX/PDF
    metrics_df["Inputs Text"] = metrics_df["Feature Set Key"].apply(renderer.as_checkbox_text)

    display_pairs = list(METRIC_DISPLAY_MAP.items())
    delta_pairs = [(f"delta_{k}", f"Δ {v}") for k, v in METRIC_DISPLAY_MAP.items()]

    for metric_key, display_name in display_pairs + delta_pairs:
        if metric_key not in metrics_df.columns:
            continue

        lower_col = f"{metric_key}_lower"
        upper_col = f"{metric_key}_upper"
        metrics_df[f"{display_name}_val"] = metrics_df[metric_key]

        if lower_col in metrics_df.columns and upper_col in metrics_df.columns:
            metrics_df[display_name] = metrics_df.apply(
                lambda row: f"{row[metric_key]:.3f} ({row[lower_col]:.3f}-{row[upper_col]:.3f})",
                axis=1,
            )
        else:
            metrics_df[display_name] = metrics_df[metric_key].apply(lambda v: f"{v:.3f}")

    return metrics_df


def _report_display_columns(renderer: FeatureSetDisplay) -> tuple[List[str], Optional[str]]:
    component_cols = renderer.component_labels if renderer.show_checkbox else []
    label_col = "Feature Set" if renderer.show_label else None
    return component_cols, label_col


def _build_report_tables(
    group: pd.DataFrame,
    renderer: FeatureSetDisplay,
) -> tuple[pd.DataFrame, pd.DataFrame, Optional[pd.DataFrame], List[str], Optional[str]]:
    component_cols, label_col = _report_display_columns(renderer)
    table_df = group.sort_values("AUROC_val", ascending=False)

    display_cols: List[str] = []
    display_cols.extend(component_cols)
    if label_col:
        display_cols.append(label_col)
    display_cols.extend(["AUROC", "AUPRC", "Brier Score", "Sensitivity", "Specificity", "F1 Score"])
    table_df_display = table_df[display_cols].copy()

    delta_candidates = ["Δ AUROC", "Δ AUPRC", "Δ Brier Score", "Δ Sensitivity", "Δ Specificity", "Δ F1 Score"]
    delta_cols: List[str] = []
    delta_cols.extend(component_cols)
    if label_col:
        delta_cols.append(label_col)
    delta_cols += [column for column in delta_candidates if column in table_df.columns]
    delta_df = table_df[delta_cols].copy() if len(delta_cols) > 1 else None
    return table_df, table_df_display, delta_df, component_cols, label_col


def generate_html_tables(metrics_df: pd.DataFrame, renderer: FeatureSetDisplay = FEATURE_SET_DISPLAY):
    """Generate styled HTML tables and manuscript-ready per-table bundles."""

    for (outcome, branch, model_name), group in metrics_df.groupby(["Outcome", "Branch", "Model"]):
        table_df, table_df_display, delta_df, component_cols, _label_col = _build_report_tables(group, renderer)

        cmap = LinearSegmentedColormap.from_list("soft_teal", ["#ffffff", "#b7dfc4"])
        cmap_r = LinearSegmentedColormap.from_list("soft_teal_r", ["#f5c2c7", "#ffffff"])

        styler = table_df_display.style
        component_indices = [
            table_df_display.columns.get_loc(col) for col in component_cols if col in table_df_display.columns
        ]
        if component_indices:
            component_styles = []
            for idx in component_indices:
                col_name = table_df_display.columns[idx]
                label_len = max(len(str(col_name)), 3)
                width_ch = min(12, max(5, label_len + 2))
                component_styles.append({"selector": f"col.col{idx}", "props": [("width", f"{width_ch}ch")]})
                component_styles.append(
                    {
                        "selector": f"th.col_heading.level0.col{idx}",
                        "props": [
                            ("width", f"{width_ch}ch"),
                            ("padding", "1px 3px"),
                            ("text-align", "center"),
                            ("white-space", "nowrap"),
                            ("font-size", "10px"),
                        ],
                    }
                )
                component_styles.append(
                    {
                        "selector": f"td.col{idx}",
                        "props": [
                            ("width", f"{width_ch}ch"),
                            ("padding", "1px 3px"),
                            ("text-align", "center"),
                            ("white-space", "nowrap"),
                            ("font-size", "10px"),
                        ],
                    }
                )
            styler.set_table_styles(component_styles, overwrite=False)

        def apply_gradient(series, val_col, color_map):
            vals = table_df.loc[series.index, val_col]
            min_v, max_v = vals.min(), vals.max()
            if max_v == min_v:
                return ["" for _ in series]
            norm = plt.Normalize(min_v, max_v)
            colors = [color_map(norm(v)) for v in vals]
            return [f"background-color: {to_hex(color)}" for color in colors]

        def highlight_best(series, val_col, mode="max"):
            vals = table_df.loc[series.index, val_col]
            is_best = vals == (vals.max() if mode == "max" else vals.min())
            return ["font-weight: bold" if value else "" for value in is_best]

        if "AUROC_val" in table_df.columns:
            styler.apply(apply_gradient, val_col="AUROC_val", color_map=cmap, subset=["AUROC"])
            styler.apply(highlight_best, val_col="AUROC_val", mode="max", subset=["AUROC"])
        if "AUPRC_val" in table_df.columns:
            styler.apply(apply_gradient, val_col="AUPRC_val", color_map=cmap, subset=["AUPRC"])
            styler.apply(highlight_best, val_col="AUPRC_val", mode="max", subset=["AUPRC"])
        if "Brier Score_val" in table_df.columns:
            styler.apply(apply_gradient, val_col="Brier Score_val", color_map=cmap_r, subset=["Brier Score"])
            styler.apply(highlight_best, val_col="Brier Score_val", mode="min", subset=["Brier Score"])
        for col in ["Sensitivity", "Specificity", "F1 Score"]:
            val_col = f"{col}_val"
            if val_col in table_df.columns:
                styler.apply(apply_gradient, val_col=val_col, color_map=cmap, subset=[col])
                styler.apply(highlight_best, val_col=val_col, mode="max", subset=[col])

        styler.hide(axis="index")

        def format_cell(value):
            if isinstance(value, str) and " (" in value:
                return value.replace(" (", "<br>(")
            return value

        styler.format(format_cell)
        styler.set_properties(**{"text-align": "center"})

        html_table = styler.to_html(table_id=f"results_{outcome}_{branch}_{model_name}", escape=False)

        delta_html = ""
        if delta_df is not None and not delta_df.empty:
            delta_styler = delta_df.style.hide(axis="index")
            delta_styler.set_properties(**{"text-align": "center"})
            delta_styler.format(format_cell)

            delta_candidates = ["Δ AUROC", "Δ AUPRC", "Δ Brier Score", "Δ Sensitivity", "Δ Specificity", "Δ F1 Score"]
            display_to_key = {
                col: f"delta_{col.split('Δ ')[-1].lower().replace(' ', '_')}" for col in delta_candidates
            }
            delta_styles = get_delta_style_info(table_df, display_to_key)

            def apply_delta_background(series, col_name):
                styles = []
                for idx in series.index:
                    bg = delta_styles.get(idx, {}).get(col_name, {}).get("bg")
                    styles.append(f"background-color: {bg}" if bg else "")
                return styles

            for col in delta_candidates:
                if col not in delta_df.columns:
                    continue
                delta_styler.apply(apply_delta_background, col_name=col, subset=[col])

            delta_html = delta_styler.to_html(table_id=f"delta_{outcome}_{branch}_{model_name}", escape=False)

        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            {NEURIPS_CSS}
            <title>Results: {outcome} ({branch}) - {model_name}</title>
        </head>
        <body>
            <h1>Model Performance Analysis</h1>
            <div class="subtitle">Outcome: {outcome} | Branch: {branch} | Model: {model_name}</div>
            {html_table}
            {('<h2>Δ vs Reference</h2>' + delta_html) if delta_html else ''}
            <div class="footer">
                Generated on {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')} | AKI Prediction Project
            </div>
        </body>
        </html>
        """

        base_path = TABLES_DIR / f"results_{outcome}_{branch}_{model_name}"
        base_path.with_suffix(".html").write_text(full_html, encoding="utf-8")
        export_table_bundle(
            base_path,
            title=f"Model Performance Analysis: {outcome} | {branch} | {model_name}",
            main_dataframe=table_df_display,
            delta_dataframe=delta_df,
            summary_text=f"Outcome: {outcome} | Branch: {branch} | Model: {model_name}",
        )

    refresh_paper_bundle(PAPER_DIR)
    print("Tables generated.")


def generate_markdown_report(metrics_df: pd.DataFrame, renderer: FeatureSetDisplay = FEATURE_SET_DISPLAY):
    """Generate a consolidated Markdown report with all model tables."""

    ordered_metrics = _order_metrics_for_reporting(metrics_df)
    sections: List[Dict[str, object]] = []
    for (outcome, branch, model_name), group in ordered_metrics.groupby(["Outcome", "Branch", "Model"], sort=False):
        _table_df, table_df_display, delta_df, _component_cols, _label_col = _build_report_tables(group, renderer)
        heading = f"Outcome: {outcome} | Branch: {branch} | Model: {model_name}"
        sections.append({"heading": heading, "dataframe": table_df_display})
        if delta_df is not None and not delta_df.empty:
            sections.append({"heading": f"{heading} | Delta vs Reference", "dataframe": delta_df})

    write_markdown_sections(REPORTS_DIR / "report.md", title="AKI Prediction Project - Results Report", sections=sections)
    refresh_paper_bundle(PAPER_DIR)

def _select_prob_column(model_group: pd.DataFrame, config: PlotConfig) -> str:
    """Choose probability column based on preference and availability."""
    if config.prefer_calibrated and "y_prob_calibrated" in model_group.columns:
        return "y_prob_calibrated"
    if not config.prefer_calibrated and "y_prob_raw" in model_group.columns:
        return "y_prob_raw"
    # Fallback to whatever exists
    return "y_prob_calibrated" if "y_prob_calibrated" in model_group.columns else "y_prob_raw"


def _filter_plot_groups(df: pd.DataFrame, config: PlotConfig) -> pd.DataFrame:
    """Apply allow/block lists to feature sets and models."""
    filtered = df.copy()
    if config.feature_set_allowlist:
        filtered = filtered[filtered['feature_set'].isin(config.feature_set_allowlist)]
    if config.feature_set_blocklist:
        filtered = filtered[~filtered['feature_set'].isin(config.feature_set_blocklist)]
    if config.model_allowlist:
        filtered = filtered[filtered['model_name'].isin(config.model_allowlist)]
    if config.model_blocklist:
        filtered = filtered[~filtered['model_name'].isin(config.model_blocklist)]
    return filtered


def _calibration_bins(
    y_true: pd.Series, y_prob: pd.Series, cfg: PlotConfig
) -> Tuple[List[float], List[float], List[int]]:
    """Compute calibration bin statistics with counts.

    Returns (mean_pred, frac_positive, count) per non-empty bin.
    """
    y_true_arr = y_true.values
    y_prob_arr = y_prob.values
    n_bins = cfg.n_bins

    if cfg.bin_strategy == "quantile":
        edges = np.quantile(y_prob_arr, np.linspace(0.0, 1.0, n_bins + 1))
    else:
        edges = np.linspace(0.0, 1.0, n_bins + 1)

    # Ensure edges are strictly increasing to avoid empty duplicates
    edges = np.unique(edges)
    mean_pred: List[float] = []
    frac_pos: List[float] = []
    counts: List[int] = []

    for i in range(len(edges) - 1):
        left, right = edges[i], edges[i + 1]
        if i == 0:
            mask = (y_prob_arr >= left) & (y_prob_arr <= right)
        else:
            mask = (y_prob_arr > left) & (y_prob_arr <= right)
        if not mask.any():
            continue
        probs_bin = y_prob_arr[mask]
        true_bin = y_true_arr[mask]
        mean_pred.append(float(np.mean(probs_bin)))
        frac_pos.append(float(np.mean(true_bin)))
        counts.append(int(len(true_bin)))

    return mean_pred, frac_pos, counts


def plot_curves(df: pd.DataFrame, config: Optional[PlotConfig] = None) -> None:
    """Generate ROC, PR, and calibration plots for each model grouping."""
    cfg = config or PlotConfig.from_env()

    df = _filter_plot_groups(df, cfg)
    if df.empty:
        logger.warning("No data left to plot after applying plot filters.")
        return

    # Set Style for NeurIPS (Serif, Academic)
    plt.style.use('seaborn-v0_8-whitegrid')
    plt.rcParams['font.family'] = 'serif'
    plt.rcParams['font.serif'] = ['Times New Roman', 'Times', 'DejaVu Serif']
    plt.rcParams['axes.labelsize'] = 12
    plt.rcParams['axes.titlesize'] = 14
    plt.rcParams['xtick.labelsize'] = 10
    plt.rcParams['ytick.labelsize'] = 10
    plt.rcParams['legend.fontsize'] = 10
    plt.rcParams['axes.spines.top'] = False
    plt.rcParams['axes.spines.right'] = False
    plt.rcParams['legend.frameon'] = True
    plt.rcParams['legend.framealpha'] = 0.9
    plt.rcParams['legend.edgecolor'] = 'white'

    tasks = list(df.groupby(['outcome', 'branch', 'model_name']))
    if not tasks:
        logger.warning("No groups found for plotting.")
        return

    def _plot_group(group_key, group_df):
        outcome, branch, model_name = group_key

        if group_df['y_true'].nunique() < 2:
            logger.warning(
                "Skipping plots for %s/%s/%s because only one class is present.",
                outcome,
                branch,
                model_name,
            )
            return

        valid_groups: Dict[str, pd.DataFrame] = {}
        for feature_set, model_group in group_df.groupby('feature_set'):
            if model_group['y_true'].nunique() < 2:
                logger.warning(
                    "Skipping feature set %s for %s/%s/%s because only one class is present.",
                    feature_set,
                    outcome,
                    branch,
                    model_name,
                )
                continue
            valid_groups[feature_set] = model_group

        if not valid_groups:
            logger.warning(
                "No feature sets with positive and negative cases for %s/%s/%s; skipping plots.",
                outcome,
                branch,
                model_name,
            )
            return

        # Helper to save and close figures consistently
        def _save_fig(fig_obj: plt.Figure, path: Path) -> None:
            save_figure_bundle(fig_obj, path, formats=("svg", "png"), close=True)

        # 1. ROC Curve
        fig, ax = plt.subplots(figsize=(8, 6), constrained_layout=True)

        auc_scores: List[Tuple[str, float]] = []
        for feature_set, model_group in valid_groups.items():
            y_true = model_group['y_true']
            prob_col = _select_prob_column(model_group, cfg)
            y_prob = model_group[prob_col]
            roc_auc = roc_auc_score(y_true, y_prob)
            auc_scores.append((feature_set, roc_auc))

        auc_scores.sort(key=lambda x: x[1], reverse=True)

        for feature_set, roc_auc in auc_scores:
            model_group = valid_groups[feature_set]
            y_true = model_group['y_true']
            prob_col = _select_prob_column(model_group, cfg)
            y_prob = model_group[prob_col]
            fpr, tpr, _ = roc_curve(y_true, y_prob)

            label_name = FEATURE_SET_MAPPING.get(feature_set, feature_set)
            ax.plot(fpr, tpr, label=f'{label_name} (AUC = {roc_auc:.3f})', linewidth=1.5)

        ax.plot([0, 1], [0, 1], 'k--', alpha=0.5, linewidth=1)
        ax.set_xlabel('False Positive Rate')
        ax.set_ylabel('True Positive Rate')
        ax.set_title(f'ROC Curves - {outcome} ({branch}) - {model_name}')
        ax.legend(loc='lower right')
        _save_fig(fig, FIGURES_DIR / f'roc_{outcome}_{branch}_{model_name}.png')

        # 2. PR Curve
        fig, ax = plt.subplots(figsize=(8, 6), constrained_layout=True)

        auprc_scores: List[Tuple[str, float]] = []
        class_balance: Dict[str, Tuple[int, int]] = {}
        for feature_set, model_group in valid_groups.items():
            y_true = model_group['y_true']
            prob_col = _select_prob_column(model_group, cfg)
            y_prob = model_group[prob_col]
            pr_auc = average_precision_score(y_true, y_prob)
            auprc_scores.append((feature_set, pr_auc))
            if cfg.show_class_balance:
                pos = int((y_true == 1).sum())
                neg = int((y_true == 0).sum())
                class_balance[feature_set] = (pos, neg)

        auprc_scores.sort(key=lambda x: x[1], reverse=True)

        prevalence = None
        if cfg.show_class_balance:
            # Use the first valid group for prevalence baseline (they share outcome/branch/model)
            for feature_set, model_group in valid_groups.items():
                y_true = model_group['y_true']
                prevalence = float((y_true == 1).mean())
                break

        if prevalence is not None:
            ax.axhline(prevalence, color='gray', linestyle='--', linewidth=1, alpha=0.7, label=f'Prevalence = {prevalence:.3f}')

        for feature_set, pr_auc in auprc_scores:
            model_group = valid_groups[feature_set]
            y_true = model_group['y_true']
            prob_col = _select_prob_column(model_group, cfg)
            y_prob = model_group[prob_col]
            prec, rec, _ = precision_recall_curve(y_true, y_prob)

            label_name = FEATURE_SET_MAPPING.get(feature_set, feature_set)
            cb_suffix = ""
            if cfg.show_class_balance and feature_set in class_balance:
                pos, neg = class_balance[feature_set]
                cb_suffix = f" (pos={pos}, neg={neg})"
            ax.step(rec, prec, where="post", label=f'{label_name} (AP = {pr_auc:.3f}){cb_suffix}', linewidth=1.5)

        ax.set_xlabel('Recall')
        ax.set_ylabel('Precision')
        ax.set_title(f'PR Curves - {outcome} ({branch}) - {model_name}')
        ax.legend(loc='lower left')
        _save_fig(fig, FIGURES_DIR / f'pr_{outcome}_{branch}_{model_name}.png')

        # 3. Calibration Curve
        fig, ax = plt.subplots(figsize=(8, 6), constrained_layout=True)
        all_probs: List[float] = []
        max_prob_seen = 0.0
        for feature_set, model_group in valid_groups.items():
            y_true = model_group['y_true']
            prob_col = _select_prob_column(model_group, cfg)
            y_prob = model_group[prob_col]
            prob_pred, prob_true, counts = _calibration_bins(y_true, y_prob, cfg)
            all_probs.extend(y_prob.tolist())
            max_prob_seen = max(max_prob_seen, y_prob.max())

            label_name = FEATURE_SET_MAPPING.get(feature_set, feature_set)
            ax.plot(prob_pred, prob_true, marker='o', label=label_name, linewidth=1, markersize=4, alpha=0.7)

            if cfg.show_bin_counts:
                for x, y, c in zip(prob_pred, prob_true, counts):
                    if c > cfg.max_count_to_annotate:
                        continue
                    ax.annotate(
                        str(c),
                        (x, y),
                        textcoords="offset points",
                        xytext=(0, 6),
                        ha='center',
                        fontsize=7,
                        color='gray',
                    )

        ax.plot([0, 1], [0, 1], 'k--', alpha=0.5, linewidth=1)
        ax.set_xlabel('Mean Predicted Probability')
        ax.set_ylabel('Fraction of Positives')
        ax.set_title(f'Calibration Curves - {outcome} ({branch}) - {model_name}')

        if max_prob_seen > 0:
            auto_xmax = min(1.0, math.ceil((max_prob_seen * 1.05) * 10) / 10.0)
            auto_xmax = max(auto_xmax, 0.1)
            auto_xmax = min(auto_xmax, 0.3)
            ax.set_xlim(0, auto_xmax)
            if cfg.show_xlim_inset:
                inset_ax = inset_axes(ax, width="35%", height="35%", loc="upper left", borderpad=1.2)
                inset_ax.plot([0, 1], [0, 1], 'k--', alpha=0.4, linewidth=1)
                inset_ax.set_xlim(0, 1)
                inset_ax.set_ylim(ax.get_ylim())
                inset_ax.set_xticks([0, 0.5, 1.0])
                inset_ax.set_yticks([])
                inset_ax.tick_params(axis='both', labelsize=7)
                inset_ax.set_title("Full range", fontsize=7)
                inset_ax.grid(False)
                for feature_set, model_group in valid_groups.items():
                    y_true = model_group['y_true']
                    prob_col = _select_prob_column(model_group, cfg)
                    y_prob = model_group[prob_col]
                    prob_pred, prob_true, _ = _calibration_bins(y_true, y_prob, cfg)
                    inset_ax.plot(prob_pred, prob_true, linewidth=1, alpha=0.4)

        if cfg.show_prob_hist and all_probs:
            rug_ax = inset_axes(
                ax,
                width="100%",
                height="22%",
                loc="lower left",
                borderpad=0,
                bbox_to_anchor=(0, -0.32, 1, 0.25),
                bbox_transform=ax.transAxes,
            )
            rug_ax.hist(all_probs, bins=min(20, cfg.n_bins * 2), color="#4c72b0", alpha=0.35, edgecolor="none")
            rug_ax.set_xlim(ax.get_xlim())
            rug_ax.set_yticks([])
            rug_ax.set_xticks(ax.get_xticks())
            rug_ax.set_xlabel('Predicted probability distribution', fontsize=8)
            rug_ax.tick_params(axis='both', labelsize=7)
            rug_ax.grid(False)

        ax.legend(loc='best')
        _save_fig(fig, FIGURES_DIR / f'calibration_{outcome}_{branch}_{model_name}.png')

    Parallel(n_jobs=cfg.n_jobs)(delayed(_plot_group)(key, group) for key, group in tasks)
    refresh_paper_bundle(PAPER_DIR)

def set_cell_background(cell, color_hex):
    """
    Set background color of a table cell in DOCX.
    color_hex: string like "#FFFFFF" or "FFFFFF"
    """
    if color_hex.startswith('#'):
        color_hex = color_hex[1:]
    
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def get_style_info(df: pd.DataFrame) -> Dict:
    """
    Calculates style information (bold, background color) for each cell.
    Returns a nested dict: {row_idx: {col_name: {'bold': bool, 'bg': hex_str}}}
    """
    style_info = {}
    
    # Colormaps
    cmap = LinearSegmentedColormap.from_list("soft_teal", ["#ffffff", "#d1e7dd"]) 
    cmap_r = LinearSegmentedColormap.from_list("soft_teal_r", ["#d1e7dd", "#ffffff"]) 
    
    # Initialize
    for idx in df.index:
        style_info[idx] = {}
        
    # Process each metric column
    metrics = ['AUROC', 'AUPRC', 'Brier Score', 'Sensitivity', 'Specificity', 'F1 Score']
    
    for col in metrics:
        val_col = f'{col}_val'
        if val_col not in df.columns:
            continue
            
        vals = df[val_col]
        min_v, max_v = vals.min(), vals.max()
        
        # Determine best
        if col == 'Brier Score':
            is_best = vals == min_v
            norm = plt.Normalize(min_v, max_v)
            current_cmap = cmap_r
        else:
            is_best = vals == max_v
            norm = plt.Normalize(min_v, max_v)
            current_cmap = cmap
            
        # Calculate styles
        for idx, val in vals.items():
            # Bold
            bold = is_best[idx]
            
            # Background
            if max_v == min_v:
                bg_color = "#ffffff"
            else:
                bg_color = to_hex(current_cmap(norm(val)))
                
            style_info[idx][col] = {'bold': bold, 'bg': bg_color}
            
    return style_info


def get_delta_style_info(df: pd.DataFrame, display_to_key: Dict[str, str]) -> Dict:
    """Background styling for delta tables; no color if CI crosses 0."""

    style_info: Dict[int, Dict[str, Dict[str, str]]] = {}
    cmap_pos = LinearSegmentedColormap.from_list("delta_pos", ["#ffffff", "#9fd4b2"])
    cmap_neg = LinearSegmentedColormap.from_list("delta_neg", ["#ffffff", "#f2b6bd"])

    for idx in df.index:
        style_info[idx] = {}

    for display_col, metric_key in display_to_key.items():
        val_col = f"delta_{metric_key.split('delta_')[-1]}"
        lower_col = f"{val_col}_lower"
        upper_col = f"{val_col}_upper"
        if val_col not in df.columns:
            continue

        vals = df[val_col]
        lowers = df[lower_col] if lower_col in df.columns else pd.Series(index=df.index, data=np.nan)
        uppers = df[upper_col] if upper_col in df.columns else pd.Series(index=df.index, data=np.nan)

        mask = (~vals.isna()) & (~lowers.isna()) & (~uppers.isna()) & ~((lowers <= 0) & (uppers >= 0))
        if not mask.any():
            continue

        abs_vals = vals.abs()[mask]
        max_abs = abs_vals.max()
        if max_abs == 0:
            continue
        norm = plt.Normalize(0, max_abs)

        for idx in df.index:
            if not mask.get(idx, False):
                continue
            val = vals.loc[idx]
            color = cmap_pos(norm(abs(val))) if val >= 0 else cmap_neg(norm(abs(val)))
            style_info[idx][display_col] = {'bg': to_hex(color)}

    return style_info


def _order_metrics_for_reporting(
    metrics_df: pd.DataFrame, model_priority: Sequence[str] = ("EBM", "XGBoost")
) -> pd.DataFrame:
    """Order metrics so preferred models render first in reports."""

    order_lookup = {name: idx for idx, name in enumerate(model_priority)}
    default_rank = len(order_lookup)
    ordered = metrics_df.copy()
    ordered["_model_rank"] = ordered["Model"].map(lambda name: order_lookup.get(name, default_rank))

    ordered = ordered.sort_values(
        by=["_model_rank", "Model", "Outcome", "Branch"],
        kind="mergesort",  # stable to preserve feature set ordering within groups
    )
    return ordered.drop(columns=["_model_rank"])

def generate_docx_report(metrics_df: pd.DataFrame, renderer: FeatureSetDisplay = FEATURE_SET_DISPLAY):
    """
    Generates a DOCX report with all tables.
    """
    ordered_metrics = _order_metrics_for_reporting(metrics_df)
    doc = Document()
    # Default to landscape orientation to give tables more horizontal space
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    doc.add_heading('AKI Prediction Project - Results Report', 0)
    
    component_cols = renderer.component_labels if renderer.show_checkbox else []
    label_source = "Feature Set" if renderer.show_label else None
    
    for (outcome, branch, model_name), group in ordered_metrics.groupby(['Outcome', 'Branch', 'Model'], sort=False):
        doc.add_heading(f'Outcome: {outcome} | Branch: {branch} | Model: {model_name}', level=1)
        
        # Sort
        table_df = group.sort_values('AUROC_val', ascending=False)
        display_cols: List[str] = []
        display_cols.extend(component_cols)
        if label_source:
            display_cols.append(label_source)
        display_cols += ['AUROC', 'AUPRC', 'Brier Score', 'Sensitivity', 'Specificity', 'F1 Score']
        
        # Get Styles
        styles = get_style_info(table_df)
        
        table_df_display = table_df[display_cols].copy()
        header_renames = {}
        if label_source:
            header_renames[label_source] = "Feature Set"
        if header_renames:
            table_df_display = table_df_display.rename(columns=header_renames)
            display_cols = [header_renames.get(col, col) for col in display_cols]
        
        # Add Table
        table = doc.add_table(rows=1, cols=len(display_cols))
        table.style = 'Table Grid'
        # Column widths: dynamic for components based on label length
        def _comp_width(col_name: str) -> float:
            # ~0.08" per character with padding, bounded
            return float(min(0.7, max(0.35, 0.08 * len(col_name) + 0.25)))
        metric_width = Inches(1.1)
        col_widths = [
            Inches(_comp_width(col)) if col in component_cols else metric_width  # keep metrics roomy
            for col in display_cols
        ]
        width_lookup = {col: width for col, width in zip(display_cols, col_widths)}
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(display_cols):
            hdr_cells[i].text = col
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].width = col_widths[i]
            
        # Rows
        for idx, row in table_df_display.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(display_cols):
                val = str(row[col])
                # Format: Value\n(CI)
                if ' (' in val:
                    val = val.replace(' (', '\n(')
                cell = row_cells[i]
                cell.text = val
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.width = col_widths[i]
                
                # Apply Styles
                if col in styles[idx]:
                    s = styles[idx][col]
                    if s['bold']:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                    set_cell_background(cell, s['bg'])
        
        doc.add_paragraph() # Spacer

        # Delta table (if available)
        delta_candidates = [
            'Δ AUROC',
            'Δ AUPRC',
            'Δ Brier Score',
            'Δ Sensitivity',
            'Δ Specificity',
            'Δ F1 Score',
        ]
        label_header = header_renames.get(label_source, label_source)
        delta_source_cols: List[str] = []
        delta_headers: List[str] = []
        # Component columns preserved as-is
        delta_source_cols.extend(component_cols)
        delta_headers.extend(component_cols)
        if label_source:
            delta_source_cols.append(label_source)
            delta_headers.append(label_header)
        delta_metric_cols = [c for c in delta_candidates if c in table_df.columns]
        delta_source_cols += delta_metric_cols
        delta_headers += delta_metric_cols
        if len(delta_headers) > 1:
            doc.add_paragraph('Δ vs Reference', style='Heading 2')
            delta_table_df = table_df[delta_source_cols].copy()
            if header_renames:
                delta_table_df = delta_table_df.rename(columns=header_renames)
            delta_headers = list(delta_table_df.columns)

            display_to_key = {col: f"delta_{col.split('Δ ')[-1].lower().replace(' ', '_')}" for col in delta_candidates}
            delta_styles = get_delta_style_info(table_df, display_to_key)

            delta_table = doc.add_table(rows=1, cols=len(delta_headers))
            delta_table.style = 'Table Grid'
            delta_col_widths = [
                width_lookup.get(col, Inches(_comp_width(col)) if col in component_cols else metric_width)
                for col in delta_headers
            ]

            hdr = delta_table.rows[0].cells
            for i, col in enumerate(delta_headers):
                hdr[i].text = col
                hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr[i].paragraphs[0].runs[0].bold = True
                hdr[i].width = delta_col_widths[i]

            for _, row in delta_table_df.iterrows():
                row_cells = delta_table.add_row().cells
                for i, col in enumerate(delta_headers):
                    val = str(row[col])
                    if ' (' in val:
                        val = val.replace(' (', '\n(')
                    row_cells[i].text = val
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[i].width = delta_col_widths[i]
                    bg = delta_styles.get(row.name, {}).get(col, {}).get('bg')
                    if bg:
                        set_cell_background(row_cells[i], bg)
            doc.add_paragraph()  # spacer
        
    doc.save(REPORTS_DIR / 'report.docx')
    refresh_paper_bundle(PAPER_DIR)
    print("DOCX report generated.")

def generate_pdf_report(metrics_df: pd.DataFrame, renderer: FeatureSetDisplay = FEATURE_SET_DISPLAY):
    """
    Generates an aggregated PDF report using Matplotlib.
    """
    ordered_metrics = _order_metrics_for_reporting(metrics_df)
    component_cols = renderer.component_labels if renderer.show_checkbox else []
    label_source = "Feature Set" if renderer.show_label else None
    
    with PdfPages(REPORTS_DIR / 'report.pdf') as pdf:
        for (outcome, branch, model_name), group in ordered_metrics.groupby(['Outcome', 'Branch', 'Model'], sort=False):
            # Sort
            table_df = group.sort_values('AUROC_val', ascending=False)
            display_cols: List[str] = []
            display_cols.extend(component_cols)
            if label_source:
                display_cols.append(label_source)
            display_cols += ['AUROC', 'AUPRC', 'Brier Score', 'Sensitivity', 'Specificity', 'F1 Score']
            
            # Get Styles
            styles = get_style_info(table_df)
            
            table_df_display = table_df[display_cols].copy()
            header_renames = {}
            if label_source:
                header_renames[label_source] = "Feature Set"
            if header_renames:
                table_df_display = table_df_display.rename(columns=header_renames)
                display_cols = [header_renames.get(c, c) for c in display_cols]
            
            # Format values for display
            display_data = []
            for _, row in table_df_display.iterrows():
                new_row = []
                for col in display_cols:
                    val = str(row[col])
                    if ' (' in val:
                        val = val.replace(' (', '\n(')
                    new_row.append(val)
                display_data.append(new_row)
            
            # Create Figure
            # Estimate height: Header + Rows * Height per row
            row_height = 0.6
            fig_height = len(table_df) * row_height + 2
            
            fig, ax = plt.subplots(figsize=(12, fig_height))
            ax.axis('off')
            ax.set_title(f"Outcome: {outcome} | Branch: {branch} | Model: {model_name}", fontsize=14, fontweight='bold', pad=20)
            
            # Create Table
            def _comp_width_fraction(col_name: str) -> float:
                # small fractions of the axis width; scale with label length
                return float(min(0.08, max(0.04, 0.01 * len(col_name) + 0.03)))
            col_widths = [_comp_width_fraction(col) if col in component_cols else 0.12 for col in display_cols]
            table = ax.table(cellText=display_data, colLabels=display_cols, loc='center', cellLoc='center', colWidths=col_widths)
            table.auto_set_font_size(False)
            table.set_fontsize(10)
            table.scale(1, 2.5) # Scale height to accommodate newlines
            
            # Style headers
            for (row, col), cell in table.get_celld().items():
                if row == 0:
                    cell.set_text_props(weight='bold')
                    cell.set_facecolor('#f0f0f0')
                else:
                    # Data rows (1-indexed in table, 0-indexed in df)
                    # We need to map back to df index
                    df_idx = table_df.index[row - 1]
                    col_name = display_cols[col]
                    
                    if col_name in styles[df_idx]:
                        s = styles[df_idx][col_name]
                        cell.set_facecolor(s['bg'])
                        if s['bold']:
                            cell.set_text_props(weight='bold')
            
            pdf.savefig(fig, bbox_inches='tight')
            plt.close()

            # Delta table as separate page if available
            delta_candidates = [
                'Δ AUROC',
                'Δ AUPRC',
                'Δ Brier Score',
                'Δ Sensitivity',
                'Δ Specificity',
                'Δ F1 Score',
            ]
            label_header = header_renames.get(label_source, label_source)
            delta_source_cols: List[str] = []
            delta_headers: List[str] = []
            delta_source_cols.extend(component_cols)
            delta_headers.extend(component_cols)
            if label_source:
                delta_source_cols.append(label_source)
                delta_headers.append(label_header)
            delta_metric_cols = [c for c in delta_candidates if c in table_df.columns]
            delta_source_cols += delta_metric_cols
            delta_headers += delta_metric_cols
            if len(delta_headers) > 1:
                delta_df = table_df[delta_source_cols].copy()
                if header_renames:
                    delta_df = delta_df.rename(columns=header_renames)
                delta_headers = list(delta_df.columns)
                display_data = []
                for _, row in delta_df.iterrows():
                    new_row = []
                    for col in delta_headers:
                        val = str(row[col])
                        if ' (' in val:
                            val = val.replace(' (', '\n(')
                        new_row.append(val)
                    display_data.append(new_row)

                row_height = 0.6
                fig_height = len(delta_df) * row_height + 2

                fig, ax = plt.subplots(figsize=(12, fig_height))
                ax.axis('off')
                ax.set_title(
                    f"Δ vs Reference - Outcome: {outcome} | Branch: {branch} | Model: {model_name}",
                    fontsize=14,
                    fontweight='bold',
                    pad=20,
                )

                delta_col_widths = [_comp_width_fraction(col) if col in component_cols else 0.12 for col in delta_headers]
                delta_table = ax.table(cellText=display_data, colLabels=delta_headers, loc='center', cellLoc='center', colWidths=delta_col_widths)
                delta_table.auto_set_font_size(False)
                delta_table.set_fontsize(10)
                delta_table.scale(1, 2.5)

                display_to_key = {col: f"delta_{col.split('Δ ')[-1].lower().replace(' ', '_')}" for col in delta_candidates}
                delta_styles = get_delta_style_info(table_df, display_to_key)

                for (row, col), cell in delta_table.get_celld().items():
                    if row == 0:
                        cell.set_text_props(weight='bold')
                        cell.set_facecolor('#f0f0f0')
                    else:
                        df_idx = delta_df.index[row - 1]
                        col_name = delta_headers[col]
                        bg = delta_styles.get(df_idx, {}).get(col_name, {}).get('bg')
                        if bg:
                            cell.set_facecolor(bg)

                pdf.savefig(fig, bbox_inches='tight')
                plt.close()
            
    refresh_paper_bundle(PAPER_DIR)
    print("PDF report generated.")

def main():
    """Deprecated entrypoint maintained for backwards compatibility."""

    raise SystemExit(
        "results_analysis now expects precomputed artifacts. "
        "Run reporting/make_report.py after generating metrics_summary.csv."
    )


if __name__ == "__main__":
    main()
