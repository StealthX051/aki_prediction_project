"""Generate descriptive statistics for the preoperative cohort.

This module loads the finalized preoperative dataset alongside the feature
lists defined in ``data_preparation.step_03_preop_prep`` to generate a
publication-ready table summarizing baseline characteristics. Provide a dataset
that still contains the raw categorical columns (i.e., before one-hot
encoding) so that baseline counts can be reported. Continuous
features are tested for normality using the Shapiro–Wilk test to decide
between mean/standard deviation or median/interquartile range summaries.
Categorical features are summarized with counts and percentages. Outputs are
saved to HTML, LaTeX, and DOCX under ``results/tables`` with human-readable labels
from :func:`reporting.display_dictionary.DisplayDictionary.feature_label`.
"""
from __future__ import annotations

import argparse
import logging
import os
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy.stats import shapiro

from data_preparation.inputs import COHORT_FILE, PREOP_PROCESSED_FILE
from data_preparation.step_03_preop_prep import CATEGORICAL_COLS, CONTINUOUS_COLS
from reporting.display_dictionary import DisplayDictionary, load_display_dictionary

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

RESULTS_DIR = Path(os.getenv("RESULTS_DIR", Path(__file__).resolve().parent.parent / "results"))
TABLES_DIR = RESULTS_DIR / "tables"
ORDERED_CATEGORICALS: Dict[str, Sequence[str]] = {
    # Preserve clinical severity ordering rather than frequency sorting.
    "asa": ("1.0", "2.0", "3.0", "4.0", "5.0", "6.0"),
}
BINARY_LABELS: Dict[str, Dict[str, str]] = {
    # Map binary codes to human-friendly labels for readability.
    "emop": {"0": "False", "1": "True"},
    "preop_htn": {"0": "False", "1": "True"},
    "preop_dm": {"0": "False", "1": "True"},
}


def _load_preop_data(path: Path) -> pd.DataFrame:
    """Load the preoperative dataset from CSV.

    Raises a clear error if the file does not exist or cannot be parsed.
    """

    if not path.exists():
        raise FileNotFoundError(f"Preoperative dataset not found at {path}")

    logger.info("Loading preoperative data from %s", path)
    return pd.read_csv(path)


def _load_optional_processed_data(path: Optional[Path]) -> Optional[pd.DataFrame]:
    """Load the outlier-handled preoperative dataset when available."""

    if path is None:
        return None

    if not path.exists():
        logger.warning("Processed dataset not found at %s; falling back to raw data for continuous features.", path)
        return None

    logger.info("Loading processed (winsorized) data from %s", path)
    return pd.read_csv(path)


def _shapiro_p_value(values: pd.Series, max_sample: int, random_state: int) -> Optional[float]:
    """Return the Shapiro–Wilk p-value for the provided series.

    SciPy's Shapiro implementation caps the supported sample size at 5,000 and
    requires at least three observations. Values beyond ``max_sample`` are
    randomly subsampled for efficiency and validity.
    """

    clean = pd.to_numeric(values, errors="coerce").dropna()
    if len(clean) < 3:
        return None

    sample = clean
    if len(clean) > max_sample:
        sample = clean.sample(max_sample, random_state=random_state)

    try:
        _, p_value = shapiro(sample)
    except ValueError as exc:
        logger.warning("Shapiro–Wilk test failed: %s", exc)
        return None

    return float(p_value)


def _summarize_continuous(
    series: pd.Series,
    *,
    alpha: float,
    max_sample: int,
    random_state: int,
) -> Tuple[str, Optional[float], str]:
    """Summarize a continuous series with normality-guided metrics."""

    numeric = pd.to_numeric(series, errors="coerce").dropna()
    if numeric.empty:
        return "No data", None, "not calculated"

    p_value = _shapiro_p_value(numeric, max_sample=max_sample, random_state=random_state)
    if p_value is not None and p_value >= alpha:
        mean = numeric.mean()
        std = numeric.std(ddof=1)
        summary = f"{mean:.2f} ± {std:.2f}"
        method = "mean ± SD"
    else:
        median = numeric.median()
        q1 = numeric.quantile(0.25)
        q3 = numeric.quantile(0.75)
        summary = f"{median:.2f} ({q1:.2f}, {q3:.2f})"
        method = "median (IQR)"

    return summary, p_value, method


def _format_category_counts(
    series: pd.Series,
    category_order: Optional[Sequence[str]] = None,
    category_labels: Optional[Dict[str, str]] = None,
) -> List[Tuple[str, int, float]]:
    """Return ordered category/count/percent tuples for a categorical column."""

    if series.empty:
        return []

    filled = series.fillna("Missing")
    total = len(filled)
    counts = filled.value_counts(dropna=False)
    ordered_items: List[Tuple[str, int]] = []

    if category_order:
        desired_order = [str(cat) for cat in category_order]
        counts.index = counts.index.astype(str)

        for label in desired_order:
            if label in counts:
                ordered_items.append((label, int(counts[label])))

        for label, count in counts.items():
            if label not in desired_order:
                ordered_items.append((str(label), int(count)))
    else:
        ordered_items = [(str(label), int(count)) for label, count in counts.items()]

    return [
        (category_labels.get(cat, cat) if category_labels else cat, count, (count / total) * 100)
        for cat, count in ordered_items
    ]


def _filter_and_validate_features(
    df: pd.DataFrame,
    *,
    continuous_features: Sequence[str],
    categorical_features: Sequence[str],
    dataset_path: Path,
) -> Tuple[List[str], List[str]]:
    """Ensure required columns exist and prevent silent omission of categoricals."""

    missing_continuous = [col for col in continuous_features if col not in df.columns]
    missing_categorical = [col for col in categorical_features if col not in df.columns]

    if missing_continuous:
        logger.warning("Missing continuous features in %s: %s", dataset_path, ", ".join(missing_continuous))

    if missing_categorical:
        message = (
            "Dataset %s is missing categorical features (%s). "
            "The PREOP_PROCESSED_FILE produced by step_03_preop_prep.py drops raw categoricals during one-hot encoding. "
            "Provide --dataset pointing to a pre-encoding dataset (e.g., COHORT_FILE) or retain categorical columns before encoding."
            % (dataset_path, ", ".join(missing_categorical))
        )
        raise ValueError(message)

    return [c for c in continuous_features if c in df.columns], [c for c in categorical_features if c in df.columns]


def _build_descriptive_table(
    df_raw: pd.DataFrame,
    df_processed: Optional[pd.DataFrame],
    *,
    continuous_features: Sequence[str],
    categorical_features: Sequence[str],
    display_dictionary: DisplayDictionary,
    alpha: float,
    max_sample: int,
    random_state: int,
) -> pd.DataFrame:
    """Compile descriptive statistics into a tidy DataFrame."""

    rows: List[Tuple[str, str, str, str]] = []

    for feature in continuous_features:
        source_series = df_processed[feature] if (df_processed is not None and feature in df_processed.columns) else df_raw[feature]
        summary, p_value, method = _summarize_continuous(
            source_series, alpha=alpha, max_sample=max_sample, random_state=random_state
        )
        label = display_dictionary.feature_label(feature, use_short=True, include_unit=True)
        p_display = "N/A" if p_value is None else f"{p_value:.3g}"
        rows.append((f"{label}, {method}", "Continuous", summary, p_display))

    for feature in categorical_features:
        label = display_dictionary.feature_label(feature, use_short=True, include_unit=True)
        category_rows = _format_category_counts(
            df_raw[feature],
            category_order=ORDERED_CATEGORICALS.get(feature),
            category_labels=BINARY_LABELS.get(feature),
        )
        rows.append((f"{label}, n (%)", "Categorical", "", ""))
        for category, count, pct in category_rows:
            rows.append((f"&nbsp;&nbsp;{category}", "", f"{count} ({pct:.1f}%)", ""))

    return pd.DataFrame(rows, columns=["Feature", "Type", "Summary", "Normality p-value"])


def _save_table(table: pd.DataFrame, output_prefix: str) -> Tuple[Path, Path, Path]:
    """Persist the table to HTML, LaTeX, and DOCX formats."""

    TABLES_DIR.mkdir(parents=True, exist_ok=True)

    html_path = TABLES_DIR / f"{output_prefix}.html"
    latex_path = TABLES_DIR / f"{output_prefix}.tex"
    docx_path = TABLES_DIR / f"{output_prefix}.docx"

    table.to_html(html_path, index=False, escape=False)
    table.to_latex(latex_path, index=False, escape=False, longtable=True)
    _save_docx(table, docx_path)

    logger.info("Saved descriptive table to %s, %s, and %s", html_path, latex_path, docx_path)
    return html_path, latex_path, docx_path


def _save_docx(table: pd.DataFrame, path: Path) -> None:
    """Write the descriptive table to a DOCX document."""

    document = Document()
    document.add_heading("Preoperative Descriptive Statistics", level=1)

    doc_table = document.add_table(rows=1 + len(table), cols=table.shape[1])
    doc_table.style = "Table Grid"

    header_cells = doc_table.rows[0].cells
    for idx, column in enumerate(table.columns):
        header_cells[idx].text = str(column)

    for row_idx, (_, row) in enumerate(table.iterrows(), start=1):
        cells = doc_table.rows[row_idx].cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)

    document.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--dataset",
        type=Path,
        default=COHORT_FILE,
        help=(
            "Path to the preoperative dataset (CSV) that still contains the raw categorical columns. "
            "If you pass PREOP_PROCESSED_FILE from step_03_preop_prep.py, reintroduce the categoricals before encoding."
        ),
    )
    parser.add_argument(
        "--processed-dataset",
        type=Path,
        default=PREOP_PROCESSED_FILE,
        help=(
            "Optional path to the outlier-handled preoperative dataset (after winsorization in step_03). "
            "Continuous features will be pulled from this file when present; categoricals always come from --dataset."
        ),
    )
    parser.add_argument(
        "--display-dictionary",
        type=Path,
        default=None,
        help="Optional path to a custom display dictionary JSON file.",
    )
    parser.add_argument(
        "--output-prefix",
        default="preop_descriptives",
        help="Filename prefix for saved tables under results/tables.",
    )
    parser.add_argument(
        "--alpha",
        type=float,
        default=0.05,
        help="Significance level for the Shapiro–Wilk normality test.",
    )
    parser.add_argument(
        "--max-normality-sample",
        type=int,
        default=5000,
        help="Maximum sample size to use for the Shapiro–Wilk test.",
    )
    parser.add_argument(
        "--random-state",
        type=int,
        default=42,
        help="Random seed used when subsampling for the Shapiro–Wilk test.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    df_raw = _load_preop_data(args.dataset)
    df_processed = _load_optional_processed_data(args.processed_dataset)
    if df_processed is not None and len(df_processed) != len(df_raw):
        logger.warning(
            "Processed dataset length (%s) does not match raw dataset length (%s); ignoring processed dataset.",
            len(df_processed),
            len(df_raw),
        )
        df_processed = None

    display_dict = load_display_dictionary(args.display_dictionary)
    total_n = len(df_raw)

    continuous_features, categorical_features = _filter_and_validate_features(
        df_raw,
        continuous_features=CONTINUOUS_COLS,
        categorical_features=CATEGORICAL_COLS,
        dataset_path=args.dataset,
    )

    descriptive_table = _build_descriptive_table(
        df_raw,
        df_processed,
        continuous_features=continuous_features,
        categorical_features=categorical_features,
        display_dictionary=display_dict,
        alpha=args.alpha,
        max_sample=args.max_normality_sample,
        random_state=args.random_state,
    )

    descriptive_table.columns = [
        "Characteristic",
        "Type",
        f"Finding (N={total_n})",
        "Normality p-value",
    ]

    if descriptive_table.empty:
        raise ValueError("No descriptive statistics generated; verify dataset columns match expectations.")

    _save_table(descriptive_table, args.output_prefix)


if __name__ == "__main__":
    main()
