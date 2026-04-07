from __future__ import annotations

from datetime import datetime, timezone
from html import unescape
import json
from pathlib import Path
from typing import Dict, Iterable, List, Mapping, Optional, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import pandas as pd


PNG_DPI = 600


def ensure_directory(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def save_figure_bundle(
    fig: plt.Figure,
    destination: Path,
    *,
    formats: Sequence[str] = ("svg", "png"),
    png_dpi: int = PNG_DPI,
    close: bool = True,
) -> List[Path]:
    """Save a figure in manuscript-ready variants."""

    base = destination.with_suffix("") if destination.suffix else destination
    saved_paths: List[Path] = []
    for fmt in formats:
        output_path = base.with_suffix(f".{fmt}")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        save_kwargs = {"bbox_inches": "tight"}
        if fmt.lower() == "png":
            save_kwargs["dpi"] = png_dpi
        fig.savefig(output_path, **save_kwargs)
        saved_paths.append(output_path)
    if close:
        plt.close(fig)
    return saved_paths


def _clean_cell(value: object) -> str:
    text = "" if pd.isna(value) else str(value)
    text = unescape(text)
    return text.replace("&nbsp;", " ").replace("\n", "<br>")


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    """Render a dataframe as a plain Markdown table without optional deps."""

    rows = [[_clean_cell(value) for value in row] for row in df.itertuples(index=False, name=None)]
    headers = [_clean_cell(column) for column in df.columns]
    widths = [len(header) for header in headers]
    for row in rows:
        for idx, value in enumerate(row):
            widths[idx] = max(widths[idx], len(value))

    def _format_row(cells: Sequence[str]) -> str:
        padded = [cell.ljust(widths[idx]) for idx, cell in enumerate(cells)]
        return "| " + " | ".join(padded) + " |"

    separator = "| " + " | ".join("-" * width for width in widths) + " |"
    lines = [_format_row(headers), separator]
    lines.extend(_format_row(row) for row in rows)
    return "\n".join(lines)


def _add_dataframe_to_doc(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, column in enumerate(df.columns):
        header_cells[idx].text = str(column)
        if header_cells[idx].paragraphs and header_cells[idx].paragraphs[0].runs:
            header_cells[idx].paragraphs[0].runs[0].bold = True
        header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row in enumerate(df.itertuples(index=False, name=None), start=1):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = _clean_cell(value).replace("<br>", "\n")
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_pdf_table_page(pdf: PdfPages, title: str, df: pd.DataFrame) -> None:
    n_rows = max(len(df), 1)
    fig_height = min(max(2.5 + 0.35 * n_rows, 4.5), 18.0)
    fig, ax = plt.subplots(figsize=(12, fig_height))
    ax.axis("off")
    ax.set_title(title, fontsize=14, fontweight="bold", pad=18)
    table = ax.table(
        cellText=[[_clean_cell(value).replace("<br>", "\n") for value in row] for row in df.itertuples(index=False, name=None)],
        colLabels=[str(column) for column in df.columns],
        cellLoc="center",
        loc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.5)
    for (row, _col), cell in table.get_celld().items():
        if row == 0:
            cell.set_text_props(weight="bold")
            cell.set_facecolor("#f2f2f2")
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def write_markdown_sections(
    path: Path,
    *,
    title: str,
    sections: Sequence[Mapping[str, object]],
) -> Path:
    """Write a Markdown document containing one or more table sections."""

    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [f"# {title}", ""]
    for section in sections:
        heading = str(section["heading"])
        df = section.get("dataframe")
        text = section.get("text")
        lines.append(f"## {heading}")
        lines.append("")
        if text:
            lines.append(str(text))
            lines.append("")
        if isinstance(df, pd.DataFrame) and not df.empty:
            lines.append(dataframe_to_markdown(df))
            lines.append("")
        elif isinstance(df, pd.DataFrame):
            lines.append("_No rows._")
            lines.append("")
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return path


def write_docx_sections(
    path: Path,
    *,
    title: str,
    sections: Sequence[Mapping[str, object]],
) -> Path:
    """Write a DOCX document containing one or more table sections."""

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    document.add_heading(title, level=0)
    for section in sections:
        document.add_heading(str(section["heading"]), level=1)
        text = section.get("text")
        if text:
            document.add_paragraph(str(text))
        df = section.get("dataframe")
        if isinstance(df, pd.DataFrame):
            _add_dataframe_to_doc(document, df)
        document.add_paragraph()
    document.save(path)
    return path


def write_pdf_sections(
    path: Path,
    *,
    title: str,
    sections: Sequence[Mapping[str, object]],
) -> Path:
    """Write a PDF document containing one or more table sections."""

    path.parent.mkdir(parents=True, exist_ok=True)
    with PdfPages(path) as pdf:
        for section in sections:
            heading = str(section["heading"])
            df = section.get("dataframe")
            text = section.get("text")
            title_text = title if heading == title else f"{title} | {heading}"
            if text and not isinstance(df, pd.DataFrame):
                fig, ax = plt.subplots(figsize=(12, 3.5))
                ax.axis("off")
                ax.set_title(title_text, fontsize=14, fontweight="bold", pad=18)
                ax.text(0.02, 0.95, str(text), va="top", ha="left", wrap=True)
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
                continue
            if isinstance(df, pd.DataFrame):
                _render_pdf_table_page(pdf, title_text, df)
    return path


def export_dataframe_bundle(
    base_path: Path,
    *,
    title: str,
    dataframe: pd.DataFrame,
    include_csv: bool = True,
) -> Dict[str, str]:
    """Export a single table to manuscript-friendly formats."""

    base = base_path.with_suffix("") if base_path.suffix else base_path
    sections = [{"heading": title, "dataframe": dataframe}]
    outputs: Dict[str, str] = {}
    if include_csv:
        csv_path = base.with_suffix(".csv")
        csv_path.parent.mkdir(parents=True, exist_ok=True)
        dataframe.to_csv(csv_path, index=False)
        outputs["csv"] = str(csv_path)
    md_path = write_markdown_sections(base.with_suffix(".md"), title=title, sections=sections)
    docx_path = write_docx_sections(base.with_suffix(".docx"), title=title, sections=sections)
    pdf_path = write_pdf_sections(base.with_suffix(".pdf"), title=title, sections=sections)
    outputs["md"] = str(md_path)
    outputs["docx"] = str(docx_path)
    outputs["pdf"] = str(pdf_path)
    return outputs


def export_table_bundle(
    base_path: Path,
    *,
    title: str,
    main_dataframe: pd.DataFrame,
    delta_dataframe: Optional[pd.DataFrame] = None,
    summary_text: Optional[str] = None,
) -> Dict[str, str]:
    """Export a results table bundle with optional delta companion table."""

    base = base_path.with_suffix("") if base_path.suffix else base_path
    sections = [{"heading": title, "dataframe": main_dataframe, "text": summary_text}]
    outputs: Dict[str, str] = {}

    main_csv = base.with_name(f"{base.name}_main").with_suffix(".csv")
    main_csv.parent.mkdir(parents=True, exist_ok=True)
    main_dataframe.to_csv(main_csv, index=False)
    outputs["csv_main"] = str(main_csv)

    if delta_dataframe is not None and not delta_dataframe.empty:
        delta_csv = base.with_name(f"{base.name}_delta").with_suffix(".csv")
        delta_dataframe.to_csv(delta_csv, index=False)
        outputs["csv_delta"] = str(delta_csv)
        sections.append({"heading": "Delta vs Reference", "dataframe": delta_dataframe})

    md_path = write_markdown_sections(base.with_suffix(".md"), title=title, sections=sections)
    docx_path = write_docx_sections(base.with_suffix(".docx"), title=title, sections=sections)
    pdf_path = write_pdf_sections(base.with_suffix(".pdf"), title=title, sections=sections)
    outputs["md"] = str(md_path)
    outputs["docx"] = str(docx_path)
    outputs["pdf"] = str(pdf_path)
    return outputs


def _asset_group(relative_path: Path) -> tuple[str, str]:
    top = relative_path.parts[0]
    if top == "reports":
        return "core_reports", "report"
    if top == "figures":
        return "manuscript_figures", "figure"
    if top == "tables":
        if relative_path.suffix == ".parquet":
            return "metadata_supporting", "supporting_data"
        return "manuscript_tables", "table"
    return "metadata_supporting", "metadata"


def _asset_id_and_variant(relative_path: Path) -> tuple[str, Optional[str]]:
    stem = relative_path.stem
    if stem.endswith("_main"):
        return stem[:-5], "main"
    if stem.endswith("_delta"):
        return stem[:-6], "delta"
    return stem, None


def _title_from_asset_id(asset_id: str, asset_type: str) -> str:
    if asset_id == "report":
        return "Consolidated Results Report"
    if asset_id == "metrics_summary":
        return "Metrics Summary"
    if asset_id == "preop_descriptives":
        return "Preoperative Descriptive Statistics"
    if asset_id == "missingness_table":
        return "Missingness Table"
    if asset_id == "cohort_flow":
        return "Cohort Flow Diagram"
    if asset_id.startswith("results_"):
        return "Results Table: " + asset_id.removeprefix("results_").replace("_", " ").title()
    if asset_type == "figure":
        return asset_id.replace("_", " ").title()
    return asset_id.replace("_", " ").title()


def refresh_paper_bundle(paper_dir: Path) -> dict:
    """Scan the paper directory and regenerate the manifest + README index."""

    paper_dir = paper_dir.resolve(strict=False)
    assets: Dict[str, dict] = {}
    for file_path in sorted(paper_dir.rglob("*")):
        if not file_path.is_file():
            continue
        relative = file_path.relative_to(paper_dir)
        if relative.as_posix() in {"manifest.json", "README.md"}:
            continue

        group, asset_type = _asset_group(relative)
        asset_id, csv_variant = _asset_id_and_variant(relative)
        entry = assets.setdefault(
            asset_id,
            {
                "id": asset_id,
                "title": _title_from_asset_id(asset_id, asset_type),
                "type": asset_type,
                "group": group,
                "paths": {},
            },
        )
        rel_text = relative.as_posix()
        if csv_variant is not None and relative.suffix == ".csv":
            entry.setdefault("csv_variants", {})[csv_variant] = rel_text
        else:
            entry["paths"][relative.suffix.lstrip(".")] = rel_text

    ordered_assets = sorted(assets.values(), key=lambda item: (item["group"], item["title"], item["id"]))
    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "paper_dir": str(paper_dir),
        "groups": {
            "core_reports": "Core reports",
            "manuscript_tables": "Manuscript-facing tables",
            "manuscript_figures": "Manuscript-facing figures",
            "metadata_supporting": "Metadata and supporting artifacts",
        },
        "assets": ordered_assets,
    }

    manifest_path = paper_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")

    lines = [
        "# Catch22 Paper Bundle",
        "",
        f"Generated: {manifest['generated_at']}",
        "",
    ]
    grouped_assets: Dict[str, List[dict]] = {
        "core_reports": [],
        "manuscript_tables": [],
        "manuscript_figures": [],
        "metadata_supporting": [],
    }
    for asset in ordered_assets:
        grouped_assets.setdefault(asset["group"], []).append(asset)

    for group_key, heading in manifest["groups"].items():
        lines.append(f"## {heading}")
        lines.append("")
        items = grouped_assets.get(group_key, [])
        if not items:
            lines.append("- None")
            lines.append("")
            continue
        for asset in items:
            format_bits = [f"{fmt}: `{path}`" for fmt, path in sorted(asset["paths"].items())]
            csv_variants = asset.get("csv_variants", {})
            format_bits.extend(f"csv-{variant}: `{path}`" for variant, path in sorted(csv_variants.items()))
            lines.append(f"- {asset['title']} ({asset['id']}): " + ", ".join(format_bits))
        lines.append("")

    readme_path = paper_dir / "README.md"
    readme_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return manifest
